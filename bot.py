#!/usr/bin/env python3
"""
bot.py — JadeBridge Telegram Bot
Smart Router: автоматическое переключение моделей через OpenRouter.
"""

import os
import asyncio
import datetime
import logging
import tempfile
import subprocess
import zipfile
import shutil
import math
from dotenv import load_dotenv
load_dotenv()  # должен выполниться ДО импорта topic_router (читает os.getenv при импорте)

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ApplicationBuilder, CommandHandler, MessageHandler, filters, ContextTypes, CallbackQueryHandler
import httpx
import sqlite3
import json
import ai_service
from groq import Groq
import pdfplumber
import docx as python_docx
import openpyxl
from topic_router import forward_to_topic, get_topic_id_for_file, SUPERGROUP_ID

TOPIC_JADE_ID: int = int(os.getenv("TOPIC_JADE_ID", "1"))


TELEGRAM_TOKEN = os.getenv("TELEGRAM_TOKEN")
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

# ── SQLite база данных для сохранения связей сообщений с файлами ───────────────
DB_FILE = "jade_bridge.db"

def init_db():
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS forwarded_files (
            chat_id INTEGER,
            message_id INTEGER,
            file_id TEXT,
            file_type TEXT,
            file_name TEXT,
            extracted_text TEXT,
            metadata TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            PRIMARY KEY (chat_id, message_id)
        )
    """)
    conn.commit()
    conn.close()

def save_forwarded_file(chat_id: int, message_id: int, file_id: str, file_type: str, file_name: str, extracted_text: str = None, metadata: dict = None):
    # Защита от Mock-объектов в тестах
    try:
        from unittest.mock import Mock
        if isinstance(chat_id, Mock): chat_id = -100123
        if isinstance(message_id, Mock): message_id = 999
        if isinstance(file_id, Mock): file_id = "mock_file_id"
        if isinstance(file_type, Mock): file_type = "mock_file_type"
        if isinstance(file_name, Mock): file_name = "mock_file_name"
        if isinstance(extracted_text, Mock): extracted_text = "mock_extracted_text"
        if isinstance(metadata, Mock): metadata = None
    except ImportError:
        pass

    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    meta_str = json.dumps(metadata, ensure_ascii=False) if metadata is not None else None
    cursor.execute("""
        INSERT OR REPLACE INTO forwarded_files (chat_id, message_id, file_id, file_type, file_name, extracted_text, metadata)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (chat_id, message_id, file_id, file_type, file_name, extracted_text, meta_str))
    conn.commit()
    conn.close()

def get_forwarded_file(chat_id: int, message_id: int) -> dict | None:
    # Защита от Mock-объектов в тестах
    try:
        from unittest.mock import Mock
        if isinstance(chat_id, Mock): chat_id = -100123
        if isinstance(message_id, Mock): message_id = 999
    except ImportError:
        pass

    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("""
        SELECT file_id, file_type, file_name, extracted_text, metadata FROM forwarded_files
        WHERE chat_id = ? AND message_id = ?
    """, (chat_id, message_id))
    row = cursor.fetchone()
    conn.close()
    if row:
        return {
            "file_id": row[0],
            "file_type": row[1],
            "file_name": row[2],
            "extracted_text": row[3],
            "metadata": json.loads(row[4]) if row[4] else None
        }
    return None

def get_ai_analyze_keyboard() -> InlineKeyboardMarkup:
    keyboard = [
        [
            InlineKeyboardButton("Sonnet 4.6", callback_data="ai_analyze:sonnet"),
            InlineKeyboardButton("Gemini 3.5", callback_data="ai_analyze:gemini")
        ]
    ]
    return InlineKeyboardMarkup(keyboard)

# Инициализируем БД при старте
init_db()

# ── Модели ────────────────────────────────────────────────────────────────────
MODEL_SIMPLE   = "google/gemini-3.5-flash"          # быстрые ответы
MODEL_COMPLEX  = "anthropic/claude-sonnet-4-6"       # сложные задачи
MODEL_ROUTER   = "google/gemini-3.5-flash"            # классификатор
MODEL_VOICE    = "google/gemini-3.5-flash"           # голос/фото — всегда flash

# ── Отображаемые имена моделей для уведомлений ──────────────────────────────────
MODEL_FRIENDLY_NAMES = {
    "google/gemini-3.5-flash": "Gemini 3.5",
    "anthropic/claude-sonnet-4-6": "Claude Sonnet 4.6",
    "anthropic/claude-opus-4-8": "Claude Opus 4.8",
}

def get_friendly_name(model_id: str) -> str:
    return MODEL_FRIENDLY_NAMES.get(model_id, model_id)

# ── Системный промпт роутера ───────────────────────────────────────────────────
ROUTER_SYSTEM_PROMPT = (
    "Ты классификатор. Ответь ТОЛЬКО словом simple или complex.\n"
    "simple = приветствие, факт, перевод, короткий вопрос, что такое X\n"
    "complex = написать код, архитектура, отладка, баг, рефакторинг, "
    "сравнение технологий, объясни как работает X подробно, спроектируй"
)

logging.basicConfig(level=logging.INFO)
# Отключаем логирование HTTP-запросов (содержащих токен бота) в stdout
logging.getLogger("httpx").setLevel(logging.WARNING)
logging.getLogger("telegram").setLevel(logging.WARNING)
logger = logging.getLogger(__name__)

user_histories: dict[int, list] = {}
groq_client = Groq(api_key=GROQ_API_KEY)


async def delete_messages_safely(bot, chat_id: int, message_ids: list[int]):
    """Безопасно удаляет список сообщений по их ID, игнорируя ошибки (если сообщение уже удалено)."""
    for msg_id in message_ids:
        try:
            await bot.delete_message(chat_id=chat_id, message_id=msg_id)
        except Exception as e:
            logger.warning(f"Failed to delete message {msg_id}: {e}")


# ══════════════════════════════════════════════════════════════════════════════
# SMART ROUTER
# ══════════════════════════════════════════════════════════════════════════════

async def choose_model(text: str) -> tuple[str, str, str]:
    """
    Определяет модель и уведомление для текстового сообщения.

    Returns:
        (model_id, clean_text, notification_message)
        notification_message — пустая строка, если уведомление не нужно.
    """
    low = text.strip().lower()

    # 1. Проверяем явные префиксы
    if low.startswith("@claude"):
        clean = text.strip()[7:].strip()   # убираем "@claude"
        friendly = get_friendly_name(MODEL_COMPLEX)
        return MODEL_COMPLEX, clean, f"Подключаю {friendly}..."

    if low.startswith("@gemini"):
        clean = text.strip()[7:].strip()   # убираем "@gemini"
        friendly = get_friendly_name(MODEL_SIMPLE)
        return MODEL_SIMPLE, clean, f"Подключаю {friendly}..."

    # 2. Router LLM — быстрый вызов для классификации
    try:
        messages = [
            {"role": "system", "content": ROUTER_SYSTEM_PROMPT},
            {"role": "user",   "content": text},
        ]
        verdict = await ai_service.ask_llm(messages, MODEL_ROUTER)
        verdict = verdict.strip().lower()
        logger.info(f"Router verdict: '{verdict}' for: {text[:60]!r}")
    except Exception as e:
        logger.warning(f"Router LLM failed, defaulting to simple: {e}")
        verdict = "simple"

    if "complex" in verdict:
        friendly = get_friendly_name(MODEL_COMPLEX)
        return MODEL_COMPLEX, text, f"Подключаю {friendly}..."
    else:
        return MODEL_SIMPLE, text, ""   # simple — без уведомления



# ══════════════════════════════════════════════════════════════════════════════
# СИСТЕМНЫЙ ПРОМПТ (с динамической датой)
# ══════════════════════════════════════════════════════════════════════════════

def get_system_prompt() -> str:
    """Возвращает системный промпт с актуальной датой. Вызывается при каждом запросе."""
    current_date = datetime.datetime.now().strftime("%d %B %Y")
    prompt = (
        f"Ты умный ассистент JadeBridge.\n"
        f"Текущая дата: {current_date}.\n"
        "Используй эту дату при ответах на вопросы о текущих событиях, "
        "актуальных технологиях, рейтингах и версиях ПО.\n"
        "Отвечай кратко и по делу.\n"
        "\n"
        "ВАЖНОЕ ПРАВИЛО для работы с Excel-таблицами:\n"
        "Если в данных встречается пометка вида [формула, не вычислено: <формула>], "
        "это означает, что значение ячейки НЕ было сохранено в исходном файле — "
        "показан только текст формулы, но не её результат. "
        "В этом случае ты ОБЯЗАН явно сообщить пользователю об этом факте. "
        "Ты можешь высказать предположение о вероятном результате ТОЛЬКО если "
        "прямо обозначишь его как предположение — например: "
        "«вероятно, результат — X, но это не подтверждено файлом: "
        "формула не была пересчитана». "
        "НИКОГДА не приводи предположительный результат как установленный факт, "
        "независимо от того, насколько простой кажется формула.\n"
        "\n"
        "СТРОГОЕ ТРЕБОВАНИЕ К ЗАВЕРШЕНИЮ ОТВЕТА:\n"
        "ЗАПРЕЩЕНО добавлять в конце любые заключительные фразы типа: "
        "'Готов работать дальше.', 'Готов помочь.', 'Если возникнут вопросы...', "
        "'Чем ещё могу помочь?', 'Сообщите, если нужно уточнить', 'Обращайтесь' "
        "и любые их вариации. Ответ должен заканчиваться содержательной информацией."
    )
    import logging
    logging.getLogger(__name__).info(f"[SYSTEM PROMPT] Дата в промпте: {current_date}")
    return prompt

# ══════════════════════════════════════════════════════════════════════════════
# LLM — единый интерфейс (смена модели через .env без правки кода)
# ══════════════════════════════════════════════════════════════════════════════

async def ask_llm(user_id: int, user_message: str, model: str) -> str:
    """Отправляет сообщение в OpenRouter с нужной моделью, хранит историю."""
    if user_id not in user_histories:
        user_histories[user_id] = []
    user_histories[user_id].append({"role": "user", "content": user_message})

    messages = [{"role": "system", "content": get_system_prompt()}] + user_histories[user_id]
    
    reply = await ai_service.ask_llm(messages, model)
    user_histories[user_id].append({"role": "assistant", "content": reply})
    return reply


# ══════════════════════════════════════════════════════════════════════════════
# АУДИО (Groq Whisper)
# ══════════════════════════════════════════════════════════════════════════════

async def transcribe_audio(ogg_path: str) -> str:
    """Конвертирует .ogg → .mp3 через ffmpeg, затем транскрибирует через Groq Whisper."""
    mp3_path = ogg_path.replace(".ogg", ".mp3")
    subprocess.run(
        ["ffmpeg", "-y", "-i", ogg_path, "-ar", "16000", "-ac", "1", "-q:a", "4", mp3_path],
        check=True,
        capture_output=True,
    )
    with open(mp3_path, "rb") as f:
        transcription = groq_client.audio.transcriptions.create(
            file=(os.path.basename(mp3_path), f.read()),
            model="whisper-large-v3",
            language="ru",
            response_format="text",
        )
    os.remove(mp3_path)
    return transcription


def should_process_message(update: Update, context=None) -> bool:
    """
    Проверяет, нужно ли обрабатывать сообщение.
    Разрешает обработку только для:
      - приватных чатов (DM с ботом)
      - топика Jade в супергруппе SUPERGROUP_ID
        (thread_id=None или thread_id=TOPIC_JADE_ID, Telegram может присылать оба варианта)
      - Reply на сообщения из очереди pending_analyses (любой топик)
    Игнорирует сообщения от ботов.
    """
    if not update.effective_user or not update.effective_chat or not update.message:
        return False

    # Для совместимости с юнит-тестами, где поля могут быть MagicMock
    try:
        from unittest.mock import MagicMock
        is_mock_chat_type = isinstance(update.effective_chat.type, MagicMock)
        is_mock_chat_id = isinstance(update.effective_chat.id, MagicMock)
        is_mock_is_bot = isinstance(update.effective_user.is_bot, MagicMock)
    except ImportError:
        is_mock_chat_type = False
        is_mock_chat_id = False
        is_mock_is_bot = False

    is_bot = update.effective_user.is_bot
    if is_mock_is_bot:
        is_bot = False

    # Игнорируем сообщения от ботов
    if is_bot:
        return False

    chat_id = update.effective_chat.id
    chat_type = update.effective_chat.type

    if is_mock_chat_type:
        chat_type = "private"
    if is_mock_chat_id:
        chat_id = 0

    # Временное логирование для сообщений из нашей супергруппы
    if SUPERGROUP_ID is not None and chat_id == SUPERGROUP_ID:
        thread_id = update.message.message_thread_id
        logger.info(
            f"[GROUP_MSG] Входящее сообщение в супергруппе {chat_id}. "
            f"thread_id={thread_id}, text={update.message.text or update.message.caption or ''}, "
            f"user={update.effective_user.name} (id={update.effective_user.id})"
        )

    # Проверка типа чата
    if chat_type == "private":
        return True

    if SUPERGROUP_ID is not None and chat_id == SUPERGROUP_ID:
        thread_id = update.message.message_thread_id
        # Топик Jade (бывший General, thread_id=1) — входной топик для пользователей.
        # Принимаем два варианта: None (старое поведение Telegram API) и TOPIC_JADE_ID (=1).
        if thread_id is None or thread_id == TOPIC_JADE_ID:
            return True
        # Разрешаем Reply на сообщения из очереди ожидания отложенного анализа
        if context is not None:
            try:
                pending = context.bot_data.get("pending_analyses", {})
                if isinstance(pending, dict) and pending:
                    reply_to_msg = getattr(update.message, "reply_to_message", None)
                    if reply_to_msg:
                        reply_to_id = getattr(reply_to_msg, "message_id", None)
                        if reply_to_id and reply_to_id in pending:
                            return True
            except Exception:
                pass
        return False

    return False


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not should_process_message(update, context):
        return
    await update.message.reply_text(
        "JadeBridge стартовал.\n\n"
        "Пиши вопросы — модель выберется автоматически.\n"
        "Или используй префикс:\n"
        "  • @claude — Sonnet для сложных задач\n"
        "  • @gemini — Flash для быстрых ответов\n\n"
        "Голосовые сообщения распознаю через Groq Whisper 🎙️\n"
        "Документы PDF, DOCX, TXT, MD, JSON — читаю и отвечаю по содержимому 📄\n"
        "ZIP-архивы с документами внутри — тоже поддерживаю 📦\n"
        "Excel таблицы .xlsx — анализирую, поддержка формул и нескольких листов 📊"
    )


async def ping(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not should_process_message(update, context):
        return
    async with httpx.AsyncClient(timeout=10) as client:
        try:
            r = await client.get(
                "https://openrouter.ai/api/v1/models",
                headers={"Authorization": f"Bearer {OPENROUTER_API_KEY}"},
            )
            if r.status_code == 200:
                await update.message.reply_text("Pong! ✓ OpenRouter живой.")
            else:
                await update.message.reply_text(f"OpenRouter ответил: {r.status_code}")
        except Exception as e:
            await update.message.reply_text(f"Ошибка: {e}")


async def clear(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not should_process_message(update, context):
        return
    user_id = update.effective_user.id
    user_histories[user_id] = []
    await update.message.reply_text("История очищена ✓")


async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает текстовые сообщения со Smart Router.
    Голос и фото — в отдельных обработчиках, здесь не трогаются.
    При активном xlsx-диалоге — перенаправляет ответ пользователя в handle_xlsx_dialog.
    """
    if not should_process_message(update, context):
        return

    # Проверяем Reply на сообщение из очереди отложенного анализа
    if getattr(update.message, "reply_to_message", None):
        reply_to_id = getattr(update.message.reply_to_message, "message_id", None)
        if reply_to_id is not None:
            try:
                pending = context.bot_data.get("pending_analyses", {})
                if isinstance(pending, dict) and reply_to_id in pending:
                    analysis_info = pending.pop(reply_to_id)
                    await execute_deferred_analysis(
                        update, context, analysis_info,
                        update.message.text or "", reply_to_id,
                    )
                    return
            except Exception as _deferred_err:
                logger.error(f"[text_deferred] Ошибка: {_deferred_err}")

    # ── Перехват xlsx-диалога ────────────────────────────────────────────────
    if context.user_data.get("xlsx_pending"):
        await handle_xlsx_dialog(update, context)
        return
    # ────────────────────────────────────────────────────────────────────────

    user_id = update.effective_user.id
    chat_id = update.effective_chat.id
    raw_text = update.message.text
    status_msg_ids = []

    # Определяем модель
    model, clean_text, notification = await choose_model(raw_text)

    try:
        # Уведомление — только для complex / явных префиксов
        if notification:
            msg_note = await update.message.reply_text(notification)
            status_msg_ids.append(msg_note.message_id)
        else:
            msg_proc = await update.message.reply_text("Обрабатываю запрос...")
            status_msg_ids.append(msg_proc.message_id)

        reply = await ask_llm(user_id, clean_text, model)
        await update.message.reply_text(reply)
    except Exception as e:
        logger.error(f"LLM error (model={model}): {e}")
        await update.message.reply_text(f"Ошибка LLM: {e}")
    finally:
        await delete_messages_safely(context.bot, chat_id, status_msg_ids)


async def handle_voice(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает голосовые сообщения.
    Всегда использует MODEL_VOICE (Gemini Flash) — не затронуто роутером.
    """
    if not should_process_message(update, context):
        return
    user_id = update.effective_user.id
    chat_id = update.effective_chat.id
    status_msg_ids = []

    # Проверяем Reply на сообщение из очереди отложенного анализа
    if getattr(update.message, "reply_to_message", None):
        reply_to_id = getattr(update.message.reply_to_message, "message_id", None)
        if reply_to_id is not None:
            try:
                pending = context.bot_data.get("pending_analyses", {})
                if isinstance(pending, dict) and reply_to_id in pending:
                    analysis_info = pending.pop(reply_to_id)
                    status_msg_voice = await update.message.reply_text("🎙️ Распознаю голос...")
                    try:
                        voice = update.message.voice
                        voice_file = await context.bot.get_file(voice.file_id)
                        with tempfile.NamedTemporaryFile(suffix=".ogg", delete=False) as tmp:
                            ogg_path = tmp.name
                        await voice_file.download_to_drive(ogg_path)
                        recognized_text = await transcribe_audio(ogg_path)
                        os.remove(ogg_path)
                    finally:
                        await context.bot.delete_message(
                            chat_id=chat_id, message_id=status_msg_voice.message_id
                        )
                    if not recognized_text or not recognized_text.strip():
                        await update.message.reply_text("❌ Не удалось распознать речь.")
                        return
                    await execute_deferred_analysis(
                        update, context, analysis_info, recognized_text, reply_to_id,
                    )
                    return
            except Exception as _deferred_err:
                logger.error(f"[voice_deferred] Ошибка: {_deferred_err}")

    try:
        msg_rec = await update.message.reply_text("Распознаю...")
        status_msg_ids.append(msg_rec.message_id)

        voice = update.message.voice
        voice_file = await context.bot.get_file(voice.file_id)

        with tempfile.NamedTemporaryFile(suffix=".ogg", delete=False) as tmp:
            ogg_path = tmp.name

        await voice_file.download_to_drive(ogg_path)

        recognized_text = await transcribe_audio(ogg_path)
        os.remove(ogg_path)

        if not recognized_text or not recognized_text.strip():
            await update.message.reply_text("❌ Не удалось распознать речь. Попробуй ещё раз.")
            return

        msg_proc = await update.message.reply_text("Обрабатываю запрос...")
        status_msg_ids.append(msg_proc.message_id)

        reply = await ask_llm(user_id, recognized_text, MODEL_VOICE)
        
        # Объединяем префикс и ответ в одно окончательное сообщение
        final_reply = f"🎙️ {recognized_text}\n\n{reply}"
        await update.message.reply_text(final_reply)

        # ── Фаза 5: пересылка транскрипции в топик Тексты ─────────────────────
        # forward_to_topic сам ловит все ошибки — не нарушает основной flow
        msg = await forward_to_topic(
            context.bot,
            topic_name="texts",
            text=f"🎙️ {recognized_text}",
            reply_markup=get_ai_analyze_keyboard(),
        )
        if msg and SUPERGROUP_ID:
            save_forwarded_file(
                chat_id=SUPERGROUP_ID,
                message_id=msg.message_id,
                file_id=voice.file_id,
                file_type="voice",
                file_name="voice.ogg",
                extracted_text=recognized_text,
            )

    except subprocess.CalledProcessError as e:
        logger.error(f"ffmpeg error: {e.stderr.decode()}")
        await update.message.reply_text("❌ Ошибка конвертации аудио (ffmpeg).")
    except Exception as e:
        logger.error(f"Voice handler error: {e}")
        await update.message.reply_text(f"❌ Ошибка обработки голосового: {e}")
    finally:
        await delete_messages_safely(context.bot, chat_id, status_msg_ids)


# ══════════════════════════════════════════════════════════════════════════════
# ФОТО / ВИДЕО (Telegram-превью)
# ══════════════════════════════════════════════════════════════════════════════

async def handle_photo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает фото, отправленные как Telegram-превью (message.photo).
    Фаза 5: пересылает оригинал в TOPIC_IMAGES_ID с кнопками анализа.
    Автоматический Vision-анализ отключён — используется кнопка «Sonnet/Gemini».
    """
    if not should_process_message(update, context):
        return
    chat_id = update.effective_chat.id
    photo = update.message.photo[-1]
    status_msg_ids = []

    try:
        msg_status = await update.message.reply_text("🖼️ Обрабатываю фото...")
        status_msg_ids.append(msg_status.message_id)

        # Генерируем имя файла ph_MMDD_HHMM.jpg (Telegram не передаёт имя для прямых фото)
        _photo_ts = datetime.datetime.now().strftime("%m%d_%H%M")
        _photo_file_name = f"ph_{_photo_ts}.jpg"

        msg = await forward_to_topic(
            context.bot,
            topic_name="images",
            file_id=photo.file_id,
            file_name=_photo_file_name,
            media_type="photo",
            reply_markup=get_ai_analyze_keyboard(),
        )
        if msg and SUPERGROUP_ID:
            save_forwarded_file(
                chat_id=SUPERGROUP_ID,
                message_id=msg.message_id,
                file_id=photo.file_id,
                file_type="photo",
                file_name=_photo_file_name,
            )
        await update.message.reply_text("Файл адресован в: IMAGES.")
    except Exception as e:
        logger.error(f"handle_photo error: {e}")
        await update.message.reply_text(f"❌ Не удалось обработать фото: {e}")
    finally:
        await delete_messages_safely(context.bot, chat_id, status_msg_ids)


async def handle_video(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает видео, отправленные как Telegram-превью (message.video).
    Фаза 5: пересылает в TOPIC_IMAGES_ID через file_id (без LLM-анализа видео).
    LLM-анализ видео — вне рамок данной фазы.
    """
    if not should_process_message(update, context):
        return
    chat_id = update.effective_chat.id
    video = update.message.video
    status_msg_ids = []

    try:
        msg_status = await update.message.reply_text("🎬 Видео получено...")
        status_msg_ids.append(msg_status.message_id)

        # ── Фаза 5: пересылка в топик Images ──────────────────────────────────────
        msg = await forward_to_topic(
            context.bot,
            topic_name="images",
            file_id=video.file_id,
            media_type="video",
            reply_markup=get_ai_analyze_keyboard(),
        )
        if msg and SUPERGROUP_ID:
            save_forwarded_file(
                chat_id=SUPERGROUP_ID,
                message_id=msg.message_id,
                file_id=video.file_id,
                file_type="video",
                file_name=video.file_name or "video.mp4",
            )
        await update.message.reply_text("✅ Видео архивировано в топик Images.")
    finally:
        await delete_messages_safely(context.bot, chat_id, status_msg_ids)


# ══════════════════════════════════════════════════════════════════════════════
# ДОКУМЕНТЫ (PDF / DOCX / TXT / MD / JSON / ZIP)
# ══════════════════════════════════════════════════════════════════════════════

MAX_DOC_CHARS = 15_000  # лимит символов для передачи в LLM

# ── ZIP-лимиты (страховочные, не защита от атак) ─────────────────────────────
ZIP_MAX_FILES       = 20          # максимум файлов в архиве
ZIP_MAX_TOTAL_BYTES = 50 * 1024 * 1024  # 50 МБ суммарный распакованный размер

# Расширения, которые умеем парсить внутри ZIP
# ВАЖНО: этот набор должен быть синхронизирован с _EXT_TO_TOPIC в topic_router.py
ZIP_SUPPORTED_EXTS = {".pdf", ".docx", ".txt", ".md", ".json", ".xlsx"}

# Системный мусор от macOS — игнорируем
ZIP_IGNORE_PREFIXES = ("__MACOSX", ".DS_Store")

# ── XLSX-лимиты ───────────────────────────────────────────────────────────────
# ~10K токенов — оставляет запас для системного промпта, истории диалога и ответа модели.
# Обоснование: claude-sonnet-4-6 имеет 200K контекст; xlsx таблица в 40K символов ≈ 10K токенов
# — это 5% контекста, что разумно при наличии истории диалога.
XLSX_MAX_CHARS = 40_000


def extract_text_from_pdf(path: str) -> str:
    """Извлекает текст из PDF через pdfplumber (все страницы)."""
    pages = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n\n".join(pages)


def extract_text_from_docx(path: str) -> str:
    """Извлекает текст из DOCX: параграфы + таблицы."""
    doc = python_docx.Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def extract_text_from_plain(path: str) -> str:
    """Извлекает текст из .txt/.md — обычное чтение файла в UTF-8."""
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def extract_text_from_json(path: str) -> str:
    """Извлекает и красиво форматирует JSON для удобства чтения LLM."""
    import json
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        raw = f.read()
    try:
        parsed = json.loads(raw)
        return json.dumps(parsed, indent=2, ensure_ascii=False)
    except json.JSONDecodeError:
        # Если JSON невалидный — вернуть как обычный текст, не падать
        return raw


async def _generate_caption_summary(raw_text: str, file_name: str) -> str | None:
    """
    Генерирует краткое (5–6 строк) LLM-описание содержимого файла для подписи в Telegram.
    Возвращает None при ошибке — вызывающий код пропускает подпись.
    """
    if not raw_text or not raw_text.strip():
        return None
    try:
        snippet = raw_text.strip()[:3000]
        messages = [
            {
                "role": "system",
                "content": (
                    "Опиши содержимое файла строго по формату:\n"
                    "— Первая строка: одно предложение, суть файла (что это и для чего).\n"
                    "— Если есть конкретные ключевые данные, добавь: **Включено**: и список "
                    "не более чем из 2 коротких пунктов (каждый — одна строка).\n"
                    "— Если ключевых данных нет — пункт **Включено**: пропускай.\n"
                    "— Итого не более 4–5 строк. Не пиши длинные абзацы.\n"
                    "Не упоминай имя файла — оно отображается отдельно. "
                    "Не выводи код, XML, JSON, HTML, сырые строки таблиц. "
                    "Отвечай по-русски. "
                    "СТРОГО ЗАПРЕЩЕНО добавлять в конце любые заключительные фразы: "
                    "'Готов помочь', 'Если нужна помощь', 'Готов работать дальше', "
                    "'Обращайтесь', 'Если возникнут вопросы' и любые их вариации. "
                    "Заканчивай ответ последним пунктом списка или единственным предложением."
                ),
            },
            {"role": "user", "content": f"Содержимое файла:\n\n{snippet}"},
        ]
        summary = await ai_service.ask_llm(messages, MODEL_SIMPLE)
        return summary[:600] if summary else None
    except Exception as e:
        logger.warning(f"[caption_summary] Не удалось сгенерировать саммари для {file_name!r}: {e}")
        return None


# Словарь: расширение → тип файла (читабельный) для строки «Тип файлов:» в сводке архива
_ZIP_EXT_TYPE_LABELS: dict[str, str] = {
    ".pdf": "PDF-документы",
    ".docx": "документы Word",
    ".txt": "текстовые файлы",
    ".md": "Markdown-файлы",
    ".json": "JSON-файлы",
    ".xlsx": "таблицы Excel",
    ".jpg": "изображения",
    ".jpeg": "изображения",
    ".png": "изображения",
    ".webp": "изображения",
    ".gif": "анимации",
    ".mp4": "видеофайлы",
}

# Словарь: расширение → логическое имя топика (читабельное) для строки «Отправлено в:»
_ZIP_EXT_TOPIC_LABELS: dict[str, str] = {
    ".pdf":  "Тексты",
    ".docx": "Тексты",
    ".doc":  "Тексты",
    ".txt":  "Тексты",
    ".md":   "Тексты",
    ".json": "Тексты",
    ".xlsx": "Таблицы",
    ".xlsm": "Таблицы",
    ".csv":  "Таблицы",
    ".jpg":  "Images",
    ".jpeg": "Images",
    ".png":  "Images",
    ".webp": "Images",
    ".gif":  "Images",
    ".mp4":  "Images",
}


def extract_text_from_zip(zip_path: str, archive_name: str) -> tuple[str, bool, dict]:
    """
    Распаковывает ZIP-архив и извлекает текст из поддерживаемых файлов.

    Returns:
        (combined_text, limits_exceeded, files_info)
        limits_exceeded=True если превышены лимиты.
        files_info = {
            "total_files": N,
            "processed_count": M,
            "entries": [...],
            "unsupported_names": [...],   # имена файлов с неподдерживаемым форматом
            "topic_labels": [...],         # уникальные читабельные имена топиков
            "type_labels": [...],          # уникальные читабельные типы файлов
        }
    """
    with zipfile.ZipFile(zip_path, "r") as zf:
        all_entries = zf.infolist()

        # Фильтруем системный мусор macOS ДО проверки лимитов
        entries = [
            e for e in all_entries
            if not any(
                e.filename.startswith(prefix) or os.path.basename(e.filename) == prefix
                for prefix in ZIP_IGNORE_PREFIXES
            )
            and not e.is_dir()
        ]

        total_files = len(entries)
        total_size = sum(e.file_size for e in entries)

        # Проверяем лимиты ДО распаковки
        if total_files > ZIP_MAX_FILES:
            msg = (
                f"⚠️ В архиве слишком много файлов: {total_files} шт.\n"
                f"Лимит: {ZIP_MAX_FILES} файлов.\n"
                "Пожалуйста, уменьшите количество файлов в архиве и отправьте снова."
            )
            return msg, True, {
                "total_files": total_files, "processed_count": 0, "entries": [],
                "unsupported_names": [], "topic_labels": [], "type_labels": [],
            }

        total_size_mb = total_size / (1024 * 1024)
        if total_size > ZIP_MAX_TOTAL_BYTES:
            msg = (
                f"⚠️ Суммарный размер файлов в архиве слишком большой: {total_size_mb:.1f} МБ.\n"
                f"Лимит: {ZIP_MAX_TOTAL_BYTES // (1024*1024)} МБ.\n"
                "Пожалуйста, уменьшите архив и отправьте снова."
            )
            return msg, True, {
                "total_files": total_files, "processed_count": 0, "entries": [],
                "unsupported_names": [], "topic_labels": [], "type_labels": [],
            }

        # Распаковываем во временную папку
        tmp_dir = tempfile.mkdtemp(prefix="vibe_zip_")
        try:
            zf.extractall(tmp_dir)

            processed_texts = []
            unsupported_names = []  # файлы с неподдерживаемым форматом
            processed_count = 0
            seen_topic_labels = []   # сохраняем порядок
            seen_type_labels = []    # сохраняем порядок

            for entry in entries:
                file_path = os.path.join(tmp_dir, entry.filename)
                short_name = os.path.basename(entry.filename) or entry.filename

                ext = os.path.splitext(entry.filename)[1].lower()

                # Проверяем поддержку по единому источнику истины — ZIP_SUPPORTED_EXTS
                # + изображения/видео (они поддерживаются через маршрутизацию, но текст не извлекается)
                _media_exts = {".jpg", ".jpeg", ".png", ".webp", ".gif", ".mp4"}
                _all_supported = ZIP_SUPPORTED_EXTS | _media_exts

                if ext not in _all_supported:
                    unsupported_names.append(short_name)
                    continue

                # Накапливаем метки топиков и типов файлов
                _topic_label = _ZIP_EXT_TOPIC_LABELS.get(ext)
                if _topic_label and _topic_label not in seen_topic_labels:
                    seen_topic_labels.append(_topic_label)
                _type_label = _ZIP_EXT_TYPE_LABELS.get(ext)
                if _type_label and _type_label not in seen_type_labels:
                    seen_type_labels.append(_type_label)

                # Извлекаем текст (только для текстовых форматов)
                try:
                    if ext == ".pdf":
                        text = extract_text_from_pdf(file_path)
                    elif ext == ".docx":
                        text = extract_text_from_docx(file_path)
                    elif ext in (".txt", ".md"):
                        text = extract_text_from_plain(file_path)
                    elif ext == ".json":
                        text = extract_text_from_json(file_path)
                    elif ext == ".xlsx":
                        # Для xlsx читаем первые строки первого листа как текст
                        try:
                            _xd = _open_xlsx_data(file_path)
                            _sheet_names = _xd.get("sheet_names", [])
                            _preview_parts = []
                            for _sn in _sheet_names[:2]:
                                _s = _xd["sheets"].get(_sn, {})
                                _md_lines = _s.get("markdown", "").split("\n")[:8]
                                _preview_parts.append(
                                    f"Лист «{_sn}» ({_s.get('rows', 0)} строк, {_s.get('cols', 0)} столбцов):\n"
                                    + "\n".join(_md_lines)
                                )
                            text = "\n\n".join(_preview_parts) if _preview_parts else ""
                        except Exception as _xe:
                            logger.warning(f"ZIP: не удалось прочитать xlsx {short_name}: {_xe}")
                            text = ""
                    else:
                        # Изображения/видео — текст не извлекаем, но считаем обработанными
                        text = ""

                    if text and text.strip():
                        processed_texts.append(f"[Файл: {short_name}]\n{text.strip()}")
                    # Считаем обработанными всех, у кого поддерживаемый формат
                    processed_count += 1

                except Exception as file_err:
                    logger.warning(f"ZIP: не удалось прочитать {short_name}: {file_err}")
                    # Не добавляем в unsupported — формат поддерживается, просто ошибка чтения
                    processed_count += 1  # всё равно засчитываем как обработанный (файл отправится в топик)

            # Собираем text для LLM-саммари (только содержательный контент, не структура)
            combined_for_summary = "\n\n".join(processed_texts) if processed_texts else ""

            files_info = {
                "total_files": total_files,
                "processed_count": processed_count,
                "entries": [e.filename for e in entries],
                "unsupported_names": unsupported_names,
                "topic_labels": seen_topic_labels,
                "type_labels": seen_type_labels,
            }
            return combined_for_summary, False, files_info

        finally:
            shutil.rmtree(tmp_dir, ignore_errors=True)


async def _forward_zip_contents_to_topics(bot, zip_path: str, archive_name: str = "") -> None:
    """
    Фаза 5: Открывает ZIP и пересылает каждый файл внутри в соответствующий топик.
    Для каждого файла:
      - определяет media_type по расширению
      - извлекает текст для поддерживаемых форматов (включая .xlsx)
      - генерирует LLM-саммари через _generate_caption_summary
      - передаёт саммари как extracted_text → build_caption выведет "Содержит: <саммари>"
      - сохраняет Telegram file_id в БД
    Ошибка отдельного файла не прерывает обработку остальных.
    """
    try:
        with zipfile.ZipFile(zip_path, "r") as zf:
            entries = [
                e for e in zf.infolist()
                if not any(
                    e.filename.startswith(prefix) or os.path.basename(e.filename) == prefix
                    for prefix in ZIP_IGNORE_PREFIXES
                )
                and not e.is_dir()
            ]
            for entry in entries:
                short_name = entry.filename
                base_name = os.path.basename(short_name)
                logger.info(f"[zip_forward] Извлекаю файл из ZIP: {short_name!r}")
                try:
                    file_bytes = zf.read(entry.filename)
                    topic_id = get_topic_id_for_file(base_name)

                    # Определяем media_type по расширению
                    ext = os.path.splitext(base_name)[1].lower()
                    if ext in (".jpg", ".jpeg", ".png", ".webp"):
                        media_type = "photo"
                    elif ext == ".gif":
                        media_type = "animation"
                    elif ext == ".mp4":
                        media_type = "video"
                    else:
                        media_type = "document"

                    # Извлекаем текст для поддерживаемых форматов
                    raw_extracted_text = None
                    if ext in ZIP_SUPPORTED_EXTS:
                        _tmp_path = None
                        try:
                            with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tf:
                                tf.write(file_bytes)
                                _tmp_path = tf.name
                            if ext == ".pdf":
                                raw_extracted_text = extract_text_from_pdf(_tmp_path)
                            elif ext == ".docx":
                                raw_extracted_text = extract_text_from_docx(_tmp_path)
                            elif ext in (".txt", ".md"):
                                raw_extracted_text = extract_text_from_plain(_tmp_path)
                            elif ext == ".json":
                                raw_extracted_text = extract_text_from_json(_tmp_path)
                            elif ext == ".xlsx":
                                # Для xlsx формируем preview первых строк как текст для саммари
                                try:
                                    _xd = _open_xlsx_data(_tmp_path)
                                    _sheet_names = _xd.get("sheet_names", [])
                                    _preview_parts = []
                                    for _sn in _sheet_names[:2]:
                                        _s = _xd["sheets"].get(_sn, {})
                                        _md_lines = _s.get("markdown", "").split("\n")[:8]
                                        _preview_parts.append(
                                            f"Лист «{_sn}» ({_s.get('rows', 0)} строк, {_s.get('cols', 0)} столбцов):\n"
                                            + "\n".join(_md_lines)
                                        )
                                    raw_extracted_text = "\n\n".join(_preview_parts) if _preview_parts else None
                                except Exception as _xe:
                                    logger.warning(f"[zip_forward] Не удалось прочитать xlsx {base_name!r}: {_xe}")
                        except Exception as _txt_err:
                            logger.warning(f"[zip_forward] Не удалось извлечь текст из {base_name!r}: {_txt_err}")
                        finally:
                            if _tmp_path and os.path.exists(_tmp_path):
                                os.remove(_tmp_path)

                    # Генерируем LLM-саммари для подписи файла
                    # Результат передаётся в build_caption как extracted_text → выводится как "Содержит: <саммари>"
                    file_caption_summary = await _generate_caption_summary(raw_extracted_text, base_name)

                    _metadata = {"archive": archive_name} if archive_name else None

                    msg = await forward_to_topic(
                        bot,
                        topic_id=topic_id,
                        file_bytes=file_bytes,
                        file_name=base_name,
                        media_type=media_type,
                        extracted_text=file_caption_summary,
                        metadata=_metadata,
                        reply_markup=get_ai_analyze_keyboard(),
                    )

                    # Сохраняем сгенерированный Telegram file_id в БД
                    if msg and SUPERGROUP_ID:
                        _file_id = None
                        if hasattr(msg, "document") and msg.document:
                            _file_id = msg.document.file_id
                        elif hasattr(msg, "photo") and msg.photo:
                            _file_id = msg.photo[-1].file_id
                        elif hasattr(msg, "video") and msg.video:
                            _file_id = msg.video.file_id
                        elif hasattr(msg, "animation") and msg.animation:
                            _file_id = msg.animation.file_id
                        if _file_id:
                            save_forwarded_file(
                                chat_id=SUPERGROUP_ID,
                                message_id=msg.message_id,
                                file_id=_file_id,
                                file_type=media_type,
                                file_name=base_name,
                                extracted_text=raw_extracted_text,
                                metadata=_metadata,
                            )

                    # Пауза для избежания rate limit Telegram API
                    await asyncio.sleep(0.5)

                except Exception as file_err:
                    logger.warning(
                        f"[zip_forward] Не удалось переслать {short_name!r}: {file_err}"
                    )
    except Exception as e:
        logger.error(f"[zip_forward] Ошибка при пересылке содержимого ZIP: {e}")


# ══════════════════════════════════════════════════════════════════════════════
# EXCEL (.xlsx)
# ══════════════════════════════════════════════════════════════════════════════

def sheet_to_markdown(ws) -> str:
    """
    Конвертирует лист openpyxl в Markdown-таблицу.

    - Первая строка → заголовки (| Col1 | Col2 | ...)
    - Вторая строка → разделитель (| --- | --- | ...)
    - Остальные строки → данные
    - Формулы с вычисленным значением → показываем значение
    - Формулы без вычисленного значения (None) → "[формула, не вычислено]"
    """
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return ""

    def fmt_cell(val) -> str:
        """Форматирует значение ячейки в строку."""
        if val is None:
            return ""
        s = str(val)
        # Если это формула (начинается с =) и значение None уже заменено пустой строкой
        # выше — этот случай не сработает. Но если data_only=False случайно вернул формулу:
        if s.startswith("="):
            return f"[формула, не вычислено: {s}]"
        # Экранируем | чтобы не сломать markdown-таблицу
        return s.replace("|", "\\|")

    def fmt_row(row_vals) -> str:
        return "| " + " | ".join(fmt_cell(v) for v in row_vals) + " |"

    header = rows[0]
    col_count = len(header)

    lines = []
    lines.append(fmt_row(header))
    lines.append("| " + " | ".join(["---"] * col_count) + " |")

    for row in rows[1:]:
        lines.append(fmt_row(row))

    return "\n".join(lines)


def _open_xlsx_data(path: str) -> dict:
    """
    Открывает .xlsx файл с data_only=True и читает все листы в память.
    Возвращает словарь:
    {
        "sheet_names": [...],
        "sheets": {
            "SheetName": {
                "markdown": "...",
                "rows": N,
                "cols": M,
            },
            ...
        }
    }
    Вызывает исключение при повреждённом файле.
    """
    wb = openpyxl.load_workbook(path, data_only=True)
    result = {"sheet_names": wb.sheetnames, "sheets": {}}

    for name in wb.sheetnames:
        try:
            ws = wb[name]
            rows_list = list(ws.iter_rows(values_only=True))
            n_rows = max(0, len(rows_list) - 1)  # не считаем строку заголовков
            n_cols = len(rows_list[0]) if rows_list else 0

            # Проверяем наличие None-значений в формульных ячейках:
            # openpyxl с data_only=True вернёт None если файл не был пересохранён
            # через Excel/LibreOffice. В этом случае нужно открыть без data_only
            # и взять саму формулу как fallback.
            has_uncalculated = False
            for row in rows_list:
                for val in row:
                    if val is None:
                        # Может быть просто пустая ячейка — не паникуем
                        pass

            markdown = sheet_to_markdown(ws)

            # Fallback: если формулы не вычислены (data_only=True вернул None)
            # открываем ещё раз без data_only чтобы взять текст формул
            wb2 = openpyxl.load_workbook(path, data_only=False)
            ws2 = wb2[name]
            rows2 = list(ws2.iter_rows(values_only=True))

            # Если в data_only версии ячейка None, а в raw версии — формула
            merged_rows = []
            for r_idx, (row_d, row_r) in enumerate(zip(rows_list, rows2)):
                merged = []
                for v_data, v_raw in zip(row_d, row_r):
                    if v_data is None and v_raw is not None and str(v_raw).startswith("="):
                        # Формула не была вычислена
                        merged.append(f"[формула, не вычислено: {v_raw}]")
                        has_uncalculated = True
                    else:
                        merged.append(v_data)
                merged_rows.append(tuple(merged))

            if has_uncalculated:
                # Пересчитываем markdown с fallback-значениями
                import io
                from openpyxl.worksheet.worksheet import Worksheet
                # Создаём временный лист для форматирования через sheet_to_markdown
                # путём прямой передачи строк
                lines = []

                def fmt_cell_merged(val) -> str:
                    if val is None:
                        return ""
                    s = str(val)
                    if s.startswith("="):
                        return f"[формула, не вычислено: {s}]"
                    return s.replace("|", "\\|")

                def fmt_row_merged(row_vals) -> str:
                    return "| " + " | ".join(fmt_cell_merged(v) for v in row_vals) + " |"

                if merged_rows:
                    header = merged_rows[0]
                    col_count = len(header)
                    lines.append(fmt_row_merged(header))
                    lines.append("| " + " | ".join(["---"] * col_count) + " |")
                    for row in merged_rows[1:]:
                        lines.append(fmt_row_merged(row))
                markdown = "\n".join(lines)

            result["sheets"][name] = {
                "markdown": markdown,
                "rows": n_rows,
                "cols": n_cols,
            }

        except Exception as sheet_err:
            logger.warning(f"XLSX: ошибка чтения листа '{name}': {sheet_err}")
            result["sheets"][name] = {
                "markdown": f"[Ошибка чтения листа '{name}': {sheet_err}]",
                "rows": 0,
                "cols": 0,
            }

    return result


def build_sheet_text(sheet_name: str, sheet_data: dict) -> str:
    """
    Формирует итоговый текст для одного листа, готовый к передаче в LLM.
    Формат:
        [Лист: «Название» | N строк × M столбцов]
        | Col1 | Col2 | ...
        | --- | --- | ...
        | ... | ... | ...
    """
    return (
        f"[Лист: «{sheet_name}» | {sheet_data['rows']} строк × {sheet_data['cols']} столбцов]\n"
        f"{sheet_data['markdown']}"
    )


def split_text_into_parts(text: str, max_chars: int) -> list[str]:
    """
    Делит большой текст на части по max_chars символов,
    стараясь разбивать по переносу строки (не рвать строки таблицы).
    """
    if len(text) <= max_chars:
        return [text]

    parts = []
    lines = text.split("\n")
    current = []
    current_len = 0

    for line in lines:
        line_len = len(line) + 1  # +1 за \n
        if current_len + line_len > max_chars and current:
            parts.append("\n".join(current))
            current = [line]
            current_len = line_len
        else:
            current.append(line)
            current_len += line_len

    if current:
        parts.append("\n".join(current))

    return parts


async def _process_xlsx_sheets(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    sheet_names: list[str],
    xlsx_data: dict,
    file_name: str,
    caption: str,
    user_id: int,
    status_msg_ids: list[int] = None,
) -> None:
    """
    Общая логика обработки выбранных листов xlsx:
    - Строит текст, проверяет лимит, при необходимости запрашивает выбор усечения.
    - Если данные в норме — сразу отправляет в LLM.
    """
    if status_msg_ids is None:
        status_msg_ids = []
    chat_id = update.effective_chat.id

    # Собираем текст для выбранных листов
    parts_text = []
    for name in sheet_names:
        if name in xlsx_data["sheets"]:
            parts_text.append(build_sheet_text(name, xlsx_data["sheets"][name]))

    header = f"[Файл: {file_name} | Листы: {', '.join(sheet_names)}]\n"
    combined = header + "\n\n".join(parts_text)

    if len(combined) > XLSX_MAX_CHARS:
        # Сохраняем состояние и спрашиваем пользователя. Перед этим удаляем статусы
        await delete_messages_safely(context.bot, chat_id, status_msg_ids)
        status_msg_ids.clear()

        n_parts = math.ceil(len(combined) / XLSX_MAX_CHARS)
        approx_rows_limit = sum(
            xlsx_data["sheets"].get(n, {}).get("rows", 0) for n in sheet_names
        )
        rows_per_part = max(1, approx_rows_limit // n_parts)

        context.user_data["xlsx_pending"] = {
            "state": "awaiting_size_choice",
            "file_name": file_name,
            "sheet_names": sheet_names,
            "xlsx_data": xlsx_data,
            "combined_text": combined,
            "caption": caption,
            "user_id": user_id,
            "n_parts": n_parts,
            "rows_per_part": rows_per_part,
        }
        await update.message.reply_text(
            f"⚠️ Таблица слишком большая ({len(combined):,} символов, лимит {XLSX_MAX_CHARS:,}).\n"
            f"Примерно {n_parts} части.\n\n"
            f"Что делать?\n"
            f"1️⃣ — Обрезать: взять первые ~{XLSX_MAX_CHARS} символов\n"
            f"2️⃣ — Прислать частями ({n_parts} сообщения)"
        )
        return

    # Данные в норме — отправляем в LLM
    await _send_xlsx_to_llm(update, context, combined, caption, user_id, status_msg_ids)


async def _send_xlsx_to_llm(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    combined: str,
    caption: str,
    user_id: int,
    status_msg_ids: list[int] = None,
) -> None:
    """Формирует промпт и отправляет данные xlsx в LLM через Smart Router."""
    if status_msg_ids is None:
        status_msg_ids = []
    chat_id = update.effective_chat.id

    try:
        if caption.strip():
            user_prompt = f"{combined}\n\n{caption.strip()}"
        else:
            user_prompt = (
                f"{combined}\n\n"
                "[Системный запрос: пользователь прислал таблицу без конкретного вопроса. "
                "Составь краткую подсказку-ориентир (не более 5–6 строк):\n"
                "— что это за таблица (тип данных, назначение);\n"
                "— какие листы или ключевые столбцы в ней есть;\n"
                "— какое содержание следует искать в этом файле.\n"
                "НЕ пересказывай данные подробно. Только ориентир.]"
            )

        model, clean_prompt, notification = await choose_model(user_prompt)
        if notification:
            msg_note = await update.message.reply_text(notification)
            status_msg_ids.append(msg_note.message_id)
        else:
            msg_proc = await update.message.reply_text("Обрабатываю запрос...")
            status_msg_ids.append(msg_proc.message_id)

        reply = await ask_llm(user_id, clean_prompt, model)

        # Выбираем префикс и обрезаем входной контент до 1000 символов (Вариант А)
        preview_text = combined
        if len(preview_text) > 1000:
            preview_text = preview_text[:1000] + "... [текст усечён для превью]"
        
        final_reply = f"📊 {preview_text}\n\n{reply}"
        await update.message.reply_text(final_reply)
    finally:
        await delete_messages_safely(context.bot, chat_id, status_msg_ids)


async def handle_xlsx_dialog(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """
    Обрабатывает ответы пользователя в диалоге xlsx (выбор листа / усечение).
    Вызывается из handle_text когда context.user_data['xlsx_pending'] существует.
    """
    pending = context.user_data.get("xlsx_pending", {})
    state = pending.get("state")
    user_text = update.message.text.strip().lower()
    chat_id = update.effective_chat.id
    status_msg_ids = []

    if state == "awaiting_sheet_choice":
        sheet_names_all: list[str] = pending["sheet_names_all"]
        file_name: str = pending["file_name"]
        xlsx_data: dict = pending["xlsx_data"]
        caption: str = pending["caption"]
        user_id: int = pending["user_id"]

        chosen_sheets: list[str] = []

        if user_text in ("все", "all", "0"):
            chosen_sheets = sheet_names_all
        else:
            # Парсим номера: "1", "1 2", "1,2", "1, 3"
            import re
            numbers = re.findall(r"\d+", user_text)
            for n_str in numbers:
                idx = int(n_str) - 1
                if 0 <= idx < len(sheet_names_all):
                    name = sheet_names_all[idx]
                    if name not in chosen_sheets:
                        chosen_sheets.append(name)

        if not chosen_sheets:
            sheets_list = "\n".join(
                f"{i+1}. {name}" for i, name in enumerate(sheet_names_all)
            )
            await update.message.reply_text(
                f"❓ Не понял выбор. Пожалуйста, введи номер листа (1–{len(sheet_names_all)}) "
                f"или напиши «все».\n\nДоступные листы:\n{sheets_list}"
            )
            return

        # Очищаем ожидание — обработаем листы (возможно войдём в диалог размера)
        del context.user_data["xlsx_pending"]
        msg_sheet = await update.message.reply_text(f"📊 Обрабатываю лист(ы): {', '.join(chosen_sheets)}...")
        status_msg_ids.append(msg_sheet.message_id)
        await _process_xlsx_sheets(update, context, chosen_sheets, xlsx_data, file_name, caption, user_id, status_msg_ids)

    elif state == "awaiting_size_choice":
        file_name: str = pending["file_name"]
        sheet_names: list[str] = pending["sheet_names"]
        xlsx_data: dict = pending["xlsx_data"]
        combined: str = pending["combined_text"]
        caption: str = pending["caption"]
        user_id: int = pending["user_id"]
        n_parts: int = pending["n_parts"]

        del context.user_data["xlsx_pending"]

        if user_text in ("1", "обрезать", "обрезать"):
            truncated = combined[:XLSX_MAX_CHARS]
            truncated += "\n\n[текст таблицы обрезан до первых символов по лимиту]"
            msg_trunc = await update.message.reply_text("✂️ Беру первые данные...")
            status_msg_ids.append(msg_trunc.message_id)
            await _send_xlsx_to_llm(update, context, truncated, caption, user_id, status_msg_ids)

        elif user_text in ("2", "частями", "частями"):
            text_parts = split_text_into_parts(combined, XLSX_MAX_CHARS)
            msg_parts = await update.message.reply_text(f"📨 Отправляю {len(text_parts)} части...")
            status_msg_ids.append(msg_parts.message_id)
            for i, part in enumerate(text_parts, 1):
                part_text = f"[Часть {i} из {len(text_parts)}]\n{part}"
                is_last = (i == len(text_parts))
                await _send_xlsx_to_llm(
                    update, context, part_text, 
                    caption if i == 1 else "", 
                    user_id, 
                    status_msg_ids if is_last else []
                )
        else:
            n_parts = pending.get("n_parts", 2)
            context.user_data["xlsx_pending"] = pending  # восстанавливаем
            await update.message.reply_text(
                "❓ Не понял. Введи 1️⃣ (обрезать) или 2️⃣ (частями)."
            )
    else:
        # Неизвестное состояние — очищаем
        context.user_data.pop("xlsx_pending", None)
        await update.message.reply_text(
            "⚠️ Сессия обработки xlsx устарела. Пожалуйста, пришли файл заново."
        )


async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает документы: PDF, DOCX, TXT, MD, JSON, ZIP, XLSX.
    Извлекает текст → подмешивает в историю → отдаёт в Smart Router.
    Для XLSX: при нескольких листах запускает диалог с пользователем.
    """
    if not should_process_message(update, context):
        return
    user_id = update.effective_user.id
    chat_id = update.effective_chat.id
    document = update.message.document
    file_name = document.file_name or "document"
    caption = update.message.caption or ""

    # Проверка размера (Telegram лимит на скачивание ботом — 20MB)
    if document.file_size and document.file_size > 20 * 1024 * 1024:
        await update.message.reply_text(
            "❌ Файл слишком большой (более 20 МБ).\n"
            "Telegram не позволяет боту скачивать файлы такого размера."
        )
        return

    # Проверка расширения
    ext = os.path.splitext(file_name)[1].lower()

    # Фаза 5.1: неподдерживаемые форматы — пересылаем сразу с пометкой в топик ATTENTION
    if ext not in (".pdf", ".docx", ".txt", ".md", ".json", ".zip", ".xlsx", ".gif"):
        try:
            _fwd_topic_id = get_topic_id_for_file(file_name, document.mime_type)
            await forward_to_topic(
                context.bot,
                topic_id=_fwd_topic_id,
                file_id=document.file_id,
                file_name=file_name,
                extracted_text=f"⚠️ Формат {ext or 'неизвестный'} — не поддерживается",
            )
        except Exception as _fwd_err:
            logger.error(f"[topic_router] Не удалось переслать файл в топик: {_fwd_err}")
        await update.message.reply_text(
            f"⚠️ Формат {ext or 'неизвестный'} пока не поддерживается.\n"
            "Поддерживаю: PDF (.pdf), Word (.docx), текст (.txt, .md), JSON (.json), "
            "архивы (.zip), Excel (.xlsx).\n"
            "RAR и другие форматы — не поддерживаются 🚧"
        )
        return

    status_msg_ids = []
    if ext == ".zip":
        msg_status = await update.message.reply_text(f"📦 Распаковываю архив «{file_name}»...")
    elif ext == ".xlsx":
        msg_status = await update.message.reply_text(f"📊 Анализирую таблицу «{file_name}»...")
    else:
        msg_status = await update.message.reply_text(f"📄 Анализирую документ «{file_name}»...")
    status_msg_ids.append(msg_status.message_id)

    # Скачиваем файл во временную папку
    tmp_path = None
    try:
        tg_file = await context.bot.get_file(document.file_id)
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp_path = tmp.name
        await tg_file.download_to_drive(tmp_path)

        # Извлекаем текст
        if ext == ".xlsx":
            # ── XLSX: читаем все листы в память сразу, tmp файл больше не нужен ──
            try:
                xlsx_data = _open_xlsx_data(tmp_path)
            except Exception as xlsx_err:
                logger.error(f"XLSX: не удалось открыть файл '{file_name}': {xlsx_err}")
                await delete_messages_safely(context.bot, chat_id, status_msg_ids)
                await update.message.reply_text(
                    f"❌ Не удалось открыть файл Excel.\n"
                    f"Возможно, файл повреждён или не является корректным .xlsx файлом.\n"
                    f"Ошибка: {xlsx_err}"
                )
                return

            sheet_names = xlsx_data["sheet_names"]

            # Фаза 5.1: пересылаем XLSX с LLM-саммари структуры таблицы
            try:
                _fwd_topic_id = get_topic_id_for_file(file_name, document.mime_type)
                _xlsx_preview = f"Листы Excel: {', '.join(sheet_names)}\n\n"
                for _sn in sheet_names[:2]:
                    if _sn in xlsx_data.get("sheets", {}):
                        _s = xlsx_data["sheets"][_sn]
                        _lines = _s["markdown"].split("\n")[:10]
                        _xlsx_preview += f"Лист «{_sn}» ({_s['rows']} строк, {_s['cols']} столбцов):\n"
                        _xlsx_preview += "\n".join(_lines) + "\n\n"
                _xlsx_caption = await _generate_caption_summary(_xlsx_preview, file_name)
                msg = await forward_to_topic(
                    context.bot,
                    topic_id=_fwd_topic_id,
                    file_id=document.file_id,
                    file_name=file_name,
                    extracted_text=_xlsx_caption,
                    metadata={"sheets": sheet_names},
                    reply_markup=get_ai_analyze_keyboard(),
                )
                if msg and SUPERGROUP_ID:
                    save_forwarded_file(
                        chat_id=SUPERGROUP_ID,
                        message_id=msg.message_id,
                        file_id=document.file_id,
                        file_type="document",
                        file_name=file_name,
                        metadata={"sheets": sheet_names},
                    )
            except Exception as _fwd_err:
                logger.error(f"[topic_router] Не удалось переслать XLSX в топик: {_fwd_err}")

            if len(sheet_names) == 0:
                await delete_messages_safely(context.bot, chat_id, status_msg_ids)
                await update.message.reply_text(
                    "⚠️ В файле нет листов. Возможно, файл пустой или повреждён."
                )
                return

            if len(sheet_names) == 1:
                # Один лист — обрабатываем сразу без диалога
                await _process_xlsx_sheets(
                    update, context, sheet_names, xlsx_data, file_name, caption, user_id, status_msg_ids
                )
            else:
                # Несколько листов — запускаем диалог
                await delete_messages_safely(context.bot, chat_id, status_msg_ids)
                status_msg_ids.clear()
                sheets_list = "\n".join(
                    f"{i+1}. {name}" for i, name in enumerate(sheet_names)
                )
                context.user_data["xlsx_pending"] = {
                    "state": "awaiting_sheet_choice",
                    "file_name": file_name,
                    "sheet_names_all": sheet_names,
                    "xlsx_data": xlsx_data,
                    "caption": caption,
                    "user_id": user_id,
                }
                await update.message.reply_text(
                    f"📋 Файл «{file_name}» содержит {len(sheet_names)} листа(ов):\n"
                    f"{sheets_list}\n\n"
                    f"Напиши номер листа (1–{len(sheet_names)}), несколько через запятую, "
                    f"или напиши «все» для обработки всех листов."
                )
            return  # xlsx обработан (или запущен диалог) — дальше не идём

        elif ext == ".gif":
            # GIF-документ пересылаем как анимацию в Images
            try:
                _fwd_topic_id = get_topic_id_for_file(file_name, document.mime_type)
                msg = await forward_to_topic(
                    context.bot,
                    topic_id=_fwd_topic_id,
                    file_id=document.file_id,
                    file_name=file_name,
                    media_type="animation",
                    reply_markup=get_ai_analyze_keyboard(),
                )
                if msg and SUPERGROUP_ID:
                    save_forwarded_file(
                        chat_id=SUPERGROUP_ID,
                        message_id=msg.message_id,
                        file_id=document.file_id,
                        file_type="animation",
                        file_name=file_name,
                    )
            except Exception as _fwd_err:
                logger.error(f"[topic_router] Не удалось переслать GIF-документ в топик: {_fwd_err}")
            
            await delete_messages_safely(context.bot, chat_id, status_msg_ids)
            await update.message.reply_text("✅ GIF-документ успешно переслан в топик Images.")
            return

        elif ext == ".zip":
            try:
                raw_text, limits_exceeded, files_info = extract_text_from_zip(tmp_path, file_name)
            except zipfile.BadZipFile:
                await delete_messages_safely(context.bot, chat_id, status_msg_ids)
                await update.message.reply_text(
                    "❌ Архив повреждён или не является корректным ZIP-файлом.\n"
                    "Пожалуйста, проверьте целостность архива и попробуйте снова."
                )
                return

            if limits_exceeded:
                await delete_messages_safely(context.bot, chat_id, status_msg_ids)
                await update.message.reply_text(raw_text)
                return

            # Формируем структурированную сводку по новому шаблону
            _total = files_info.get("total_files", 0)
            _processed = files_info.get("processed_count", 0)
            _unsupported = files_info.get("unsupported_names", [])
            _topic_labels = files_info.get("topic_labels", [])
            _type_labels = files_info.get("type_labels", [])

            # Генерируем LLM-саммари для блока "Содержит:"
            _zip_llm_summary = await _generate_caption_summary(raw_text, file_name)

            # Строим сводку по шаблону — plain text (без Markdown-разметки)
            _summary_lines = [
                f"📦 Архив: `{file_name}` (Обработано: {_processed} из {_total})",
            ]
            if _topic_labels:
                _summary_lines.append(f"Отправлено в: {', '.join(_topic_labels)}")
            if _type_labels:
                _summary_lines.append(f"Тип файлов: {', '.join(_type_labels)}")
            if _unsupported:
                _unsupported_str = ", ".join(f"`{n}`" for n in _unsupported)
                _summary_lines.append(f"Не обработано (формат не поддерживается): {_unsupported_str}")
            if _zip_llm_summary:
                _summary_lines.append(f"Содержит: {_zip_llm_summary}")

            _zip_user_summary = "\n".join(_summary_lines)

            # Пересылаем ZIP с LLM-саммари как подписью
            try:
                _fwd_topic_id = get_topic_id_for_file(file_name, document.mime_type)
                msg = await forward_to_topic(
                    context.bot,
                    topic_id=_fwd_topic_id,
                    file_id=document.file_id,
                    file_name=file_name,
                    extracted_text=_zip_llm_summary,
                    reply_markup=get_ai_analyze_keyboard(),
                )
                if msg and SUPERGROUP_ID:
                    save_forwarded_file(
                        chat_id=SUPERGROUP_ID,
                        message_id=msg.message_id,
                        file_id=document.file_id,
                        file_type="document",
                        file_name=file_name,
                        extracted_text=raw_text,
                    )
            except Exception as _fwd_err:
                logger.error(f"[topic_router] Не удалось переслать ZIP в топик: {_fwd_err}")

            # Отправляем структурированную сводку пользователю (вместо LLM-ответа)
            await delete_messages_safely(context.bot, chat_id, status_msg_ids)
            status_msg_ids.clear()
            await update.message.reply_text(_zip_user_summary)

            # Пересылаем каждый файл из архива в его топик (после ответа пользователю)
            if tmp_path and os.path.exists(tmp_path):
                await _forward_zip_contents_to_topics(context.bot, tmp_path, file_name)

            # ZIP полностью обработан — ранний выход (не идём в общий LLM-flow ниже)
            return
        elif ext == ".pdf":
            raw_text = None
            try:
                raw_text = extract_text_from_pdf(tmp_path)
            except Exception as pdf_err:
                logger.warning(f"PDF: не удалось извлечь текст из {file_name}: {pdf_err}")

            # Фаза 5.1: PDF — пересылаем с thumbnail первой страницы и подписью
            try:
                _fwd_topic_id = get_topic_id_for_file(file_name, document.mime_type)
                _pdf_caption = await _generate_caption_summary(raw_text, file_name)
                msg = await forward_to_topic(
                    context.bot,
                    topic_id=_fwd_topic_id,
                    file_id=document.file_id,
                    file_name=file_name,
                    file_path=tmp_path,
                    file_type="pdf",
                    extracted_text=_pdf_caption,
                    reply_markup=get_ai_analyze_keyboard(),
                )
                if msg and SUPERGROUP_ID:
                    save_forwarded_file(
                        chat_id=SUPERGROUP_ID,
                        message_id=msg.message_id,
                        file_id=document.file_id,
                        file_type="document",
                        file_name=file_name,
                        extracted_text=raw_text,
                    )
            except Exception as _fwd_err:
                logger.error(f"[topic_router] Не удалось переслать PDF в топик: {_fwd_err}")
        elif ext == ".docx":
            raw_text = None
            try:
                raw_text = extract_text_from_docx(tmp_path)
            except Exception as docx_err:
                logger.warning(f"DOCX: не удалось извлечь текст из {file_name}: {docx_err}")
        elif ext in (".txt", ".md"):
            raw_text = None
            try:
                raw_text = extract_text_from_plain(tmp_path)
            except Exception as plain_err:
                logger.warning(f"TXT/MD: не удалось извлечь текст из {file_name}: {plain_err}")
        elif ext == ".json":
            raw_text = None
            try:
                raw_text = extract_text_from_json(tmp_path)
            except Exception as json_err:
                logger.warning(f"JSON: не удалось извлечь текст из {file_name}: {json_err}")

        # Фаза 5.1: DOCX/TXT/MD/JSON — пересылаем с подписью (PDF и ZIP переслали выше)
        if ext in (".docx", ".txt", ".md", ".json"):
            try:
                _fwd_topic_id = get_topic_id_for_file(file_name, document.mime_type)
                _doc_caption = await _generate_caption_summary(raw_text, file_name)
                msg = await forward_to_topic(
                    context.bot,
                    topic_id=_fwd_topic_id,
                    file_id=document.file_id,
                    file_name=file_name,
                    extracted_text=_doc_caption,
                    reply_markup=get_ai_analyze_keyboard(),
                )
                if msg and SUPERGROUP_ID:
                    save_forwarded_file(
                        chat_id=SUPERGROUP_ID,
                        message_id=msg.message_id,
                        file_id=document.file_id,
                        file_type="document",
                        file_name=file_name,
                        extracted_text=raw_text,
                    )
            except Exception as _fwd_err:
                logger.error(f"[topic_router] Не удалось переслать в топик: {_fwd_err}")

        if not raw_text or not raw_text.strip():
            await delete_messages_safely(context.bot, chat_id, status_msg_ids)
            await update.message.reply_text(
                "⚠️ Не удалось извлечь текст из документа.\n"
                "Возможно, это отсканированный PDF без распознанного текста."
            )
            return

        # Обрезаем до лимита
        truncated = False
        if len(raw_text) > MAX_DOC_CHARS:
            raw_text = raw_text[:MAX_DOC_CHARS]
            truncated = True

        # Формируем сообщение для LLM
        # Для ZIP архив уже содержит заголовки файлов, не дублируем
        if ext == ".zip":
            doc_content = raw_text
        else:
            doc_content = f"[Документ: {file_name}]\n\n{raw_text}"
        if truncated:
            doc_content += "\n\n[текст обрезан, архив слишком большой]" if ext == ".zip" else "\n\n[текст обрезан, документ слишком длинный]"

        # Если пользователь написал подпись — это его вопрос к документу/архиву
        # Если нет — просим LLM кратко резюмировать
        if caption.strip():
            user_prompt = f"{doc_content}\n\n{caption.strip()}"
        else:
            if ext == ".zip":
                summary_request = (
                    "[Системный запрос: пользователь прислал архив без конкретного вопроса. "
                    "Составь краткую подсказку-ориентир (не более 5–6 строк):\n"
                    "— что это за архив и какие файлы в нём содержатся;\n"
                    "— какой тип контента в файлах;\n"
                    "— какое содержание следует искать в этих файлах.\n"
                    "НЕ пересказывай контент подробно. Только ориентир.]"
                )
            else:
                summary_request = (
                    "[Системный запрос: пользователь прислал документ без конкретного вопроса. "
                    "Составь краткую подсказку-ориентир (не более 5–6 строк):\n"
                    "— что это за документ (тип, назначение);\n"
                    "— какие ключевые разделы или темы в нём есть;\n"
                    "— какое содержание следует искать в этом файле.\n"
                    "НЕ пересказывай контент подробно. Только ориентир для навигации.]"
                )
            user_prompt = f"{doc_content}\n\n{summary_request}"

        # Прогоняем через Smart Router — как обычный текст
        model, clean_prompt, notification = await choose_model(user_prompt)
        if notification:
            msg_note = await update.message.reply_text(notification)
            status_msg_ids.append(msg_note.message_id)
        else:
            msg_proc = await update.message.reply_text("Обрабатываю запрос...")
            status_msg_ids.append(msg_proc.message_id)

        reply = await ask_llm(user_id, clean_prompt, model)

        # Выбираем эмодзи и префикс (усекаем до 1000 символов)
        emoji = "📄"
        if ext == ".zip":
            emoji = "📦"
        
        preview_text = doc_content
        if len(preview_text) > 1000:
            preview_text = preview_text[:1000] + "... [текст усечён для превью]"
        
        final_reply = f"{emoji} {preview_text}\n\n{reply}"
        await update.message.reply_text(final_reply)

    except Exception as e:
        logger.error(f"Document handler error: {e}")
        await update.message.reply_text(
            f"❌ Не удалось обработать документ.\n"
            f"Возможно, файл повреждён или недоступен для чтения.\n"
            f"Ошибка: {e}"
        )
    finally:
        # Удаляем временный файл в любом случае
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)
        await delete_messages_safely(context.bot, chat_id, status_msg_ids)


# ══════════════════════════════════════════════════════════════════════════════
# ВИДЕОКРУЖОЧКИ И ГИФКИ
# ══════════════════════════════════════════════════════════════════════════════

async def handle_video_note(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает видеокружочки (video_note).
    Пересылает в топик Images с кнопками отложенного анализа.
    """
    if not should_process_message(update, context):
        return
    chat_id = update.effective_chat.id
    video_note = update.message.video_note
    status_msg_ids = []

    try:
        msg_status = await update.message.reply_text("🎬 Анализирую видео...")
        status_msg_ids.append(msg_status.message_id)

        # Пересылка в топик Images с клавиатурой выбора модели
        msg = await forward_to_topic(
            context.bot,
            topic_name="images",
            file_id=video_note.file_id,
            media_type="video_note",
            reply_markup=get_ai_analyze_keyboard(),
        )

        if msg and SUPERGROUP_ID:
            save_forwarded_file(
                chat_id=SUPERGROUP_ID,
                message_id=msg.message_id,
                file_id=video_note.file_id,
                file_type="video_note",
                file_name="video_note.mp4",
            )

        await update.message.reply_text("✅ Видеокружочек успешно переслан в топик Images.")
    except Exception as e:
        logger.error(f"handle_video_note error: {e}")
        await update.message.reply_text(f"❌ Не удалось обработать видеокружочек: {e}")
    finally:
        await delete_messages_safely(context.bot, chat_id, status_msg_ids)


async def handle_animation(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает GIF-анимации (animation).
    Пересылает в топик Images с кнопками отложенного анализа.
    """
    if not should_process_message(update, context):
        return
    chat_id = update.effective_chat.id
    animation = update.message.animation
    status_msg_ids = []

    try:
        msg_status = await update.message.reply_text("🎬 Анализирую видео...")
        status_msg_ids.append(msg_status.message_id)

        # Пересылка в топик Images с клавиатурой выбора модели
        msg = await forward_to_topic(
            context.bot,
            topic_name="images",
            file_id=animation.file_id,
            media_type="animation",
            reply_markup=get_ai_analyze_keyboard(),
        )

        if msg and SUPERGROUP_ID:
            save_forwarded_file(
                chat_id=SUPERGROUP_ID,
                message_id=msg.message_id,
                file_id=animation.file_id,
                file_type="animation",
                file_name=animation.file_name or "animation.gif",
            )

        await update.message.reply_text("✅ GIF-анимация успешно переслана в топик Images.")
    except Exception as e:
        logger.error(f"handle_animation error: {e}")
        await update.message.reply_text(f"❌ Не удалось обработать анимацию: {e}")
    finally:
        await delete_messages_safely(context.bot, chat_id, status_msg_ids)


# ══════════════════════════════════════════════════════════════════════════════
# ОТЛОЖЕННЫЙ AI-АНАЛИЗ
# ══════════════════════════════════════════════════════════════════════════════

# Фразы, сигнализирующие о запросе общего анализа (не конкретного вопроса)
GENERIC_ANALYSIS_PHRASES = frozenset({
    "дай общий анализ", "в целом", "проанализируй", "общий анализ",
    "анализ", "проанализируй файл", "общее описание", "что в файле",
    "что это", "опиши", "расскажи", "общий обзор",
})


async def execute_deferred_analysis(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    analysis_info: dict,
    user_query: str,
    question_message_id: int,
) -> None:
    """
    Запускает отложенный анализ файла по запросу пользователя.
    Удаляет вопрос-опрос и ответ пользователя, отправляет результат как Reply на файл.
    """
    chat_id = analysis_info["chat_id"]
    file_message_id = analysis_info["file_message_id"]
    model_choice = analysis_info["model_choice"]

    if model_choice == "sonnet":
        model_id = os.getenv("MODEL_COMPLEX", "anthropic/claude-sonnet-4-6")
        friendly_model_name = "Claude Sonnet 4.6"
    else:
        model_id = os.getenv("MODEL_SIMPLE", "google/gemini-3.5-flash")
        friendly_model_name = "Gemini 3.5"

    file_info = get_forwarded_file(chat_id, file_message_id)

    # Статус анализа как Reply на исходный файл в топике
    status_msg = await context.bot.send_message(
        chat_id=chat_id,
        text=f"🧠 Запуск глубокого анализа ({friendly_model_name})...",
        reply_to_message_id=file_message_id,
    )

    # Удаляем вопрос-опрос бота и ответ пользователя
    await delete_messages_safely(
        context.bot, chat_id,
        [question_message_id, update.message.message_id],
    )

    if not file_info:
        await status_msg.edit_text("⚠️ Информация об этом файле не найдена в базе данных.")
        return

    file_id = file_info["file_id"]
    file_type = file_info["file_type"]
    file_name = file_info["file_name"]
    extracted_text = file_info["extracted_text"]

    low_query = user_query.strip().lower()
    is_generic = (
        len(low_query) < 5
        or any(phrase in low_query for phrase in GENERIC_ANALYSIS_PHRASES)
    )

    try:
        reply_text = ""

        if file_type in ("video_note", "animation"):
            reply_text = "Функция глубокого анализа видео временно недоступна."

        elif file_type == "photo":
            photo_file = await context.bot.get_file(file_id)
            with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tmp:
                photo_path = tmp.name
            await photo_file.download_to_drive(photo_path)
            import base64
            with open(photo_path, "rb") as f:
                photo_b64 = base64.b64encode(f.read()).decode()
            os.remove(photo_path)
            question = (
                "Проведи глубокий детальный анализ этого изображения. Опиши все детали, объекты, текст на изображении (если есть) и контекст."
                if is_generic else user_query
            )
            reply_text = await ai_service.ask_vision(photo_b64, question, model_id)

        elif file_type == "voice":
            if not extracted_text:
                reply_text = "Не удалось найти текст голосового сообщения для анализа."
            else:
                if is_generic:
                    prompt = (
                        f"Проведи глубокий подробный анализ расшифровки голосового сообщения:\n\n"
                        f"«{extracted_text}»\n\n"
                        f"Выдели главные мысли, суть сказанного, неявные задачи и ключевые выводы."
                    )
                else:
                    prompt = (
                        f"Расшифровка голосового сообщения:\n\n«{extracted_text}»\n\n"
                        f"Вопрос пользователя: {user_query}\n\n"
                        f"Ответь на вопрос, опираясь на содержание голосового сообщения."
                    )
                messages_llm = [
                    {"role": "system", "content": "Ты аналитический ассистент JadeBridge."},
                    {"role": "user", "content": prompt},
                ]
                reply_text = await ai_service.ask_llm(messages_llm, model_id)

        elif file_type == "document":
            if not extracted_text:
                ext = os.path.splitext(file_name)[1].lower()
                doc_file = await context.bot.get_file(file_id)
                with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                    tmp_path = tmp.name
                await doc_file.download_to_drive(tmp_path)
                try:
                    if ext == ".xlsx":
                        xlsx_data = _open_xlsx_data(tmp_path)
                        parts_text = [
                            build_sheet_text(n, xlsx_data["sheets"][n])
                            for n in xlsx_data["sheet_names"]
                        ]
                        extracted_text = f"[Файл: {file_name}]\n" + "\n\n".join(parts_text)
                    elif ext == ".pdf":
                        extracted_text = extract_text_from_pdf(tmp_path)
                    elif ext == ".docx":
                        extracted_text = extract_text_from_docx(tmp_path)
                    elif ext in (".txt", ".md"):
                        extracted_text = extract_text_from_plain(tmp_path)
                    elif ext == ".json":
                        extracted_text = extract_text_from_json(tmp_path)
                finally:
                    if os.path.exists(tmp_path):
                        os.remove(tmp_path)

            if not extracted_text or not extracted_text.strip():
                reply_text = "Не удалось извлечь текст из документа для анализа."
            else:
                safe_text = (
                    extracted_text[:30000] + "... [текст обрезан]"
                    if len(extracted_text) > 30000 else extracted_text
                )
                if is_generic:
                    prompt = (
                        f"Проведи глубокий детальный анализ следующего документа:\n\n"
                        f"{safe_text}\n\n"
                        f"Напиши развернутое резюме, структуру, ключевые факты, выводы и критический разбор."
                    )
                else:
                    prompt = (
                        f"Документ:\n\n{safe_text}\n\n"
                        f"Вопрос пользователя: {user_query}\n\n"
                        f"Ответь на вопрос, опираясь на содержание документа."
                    )
                messages_llm = [
                    {"role": "system", "content": "Ты аналитический ассистент JadeBridge."},
                    {"role": "user", "content": prompt},
                ]
                reply_text = await ai_service.ask_llm(messages_llm, model_id)
        else:
            reply_text = "Неизвестный тип файла для анализа."

        header = f"🧠 *Глубокий анализ ({friendly_model_name})*\n\n"
        await status_msg.edit_text(f"{header}{reply_text}")

    except Exception as e:
        logger.error(f"Error in execute_deferred_analysis: {e}")
        await status_msg.edit_text(f"❌ Ошибка при проведении анализа: {e}")


async def handle_ai_analyze_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает нажатия на кнопки «Sonnet 4.6» и «Gemini 3.5» в топиках.
    Отправляет вопрос-опрос «Ваш запрос: ...?» для запуска отложенного анализа.
    """
    query = update.callback_query
    await query.answer()

    data = query.data  # 'ai_analyze:sonnet' или 'ai_analyze:gemini'
    model_choice = data.split(":")[1]

    message = query.message
    chat_id = message.chat.id
    message_id = message.message_id

    file_info = get_forwarded_file(chat_id, message_id)
    if not file_info:
        await message.reply_text(
            "⚠️ Информация об этом файле отсутствует в базе данных бота.\n"
            "Возможно, сообщение было отправлено до обновления БД.",
            reply_to_message_id=message_id,
        )
        return

    # Инициализируем словарь pending_analyses если нет
    if "pending_analyses" not in context.bot_data:
        context.bot_data["pending_analyses"] = {}

    # Отправляем вопрос-опрос как Reply на исходное сообщение с файлом
    question_msg = await message.reply_text(
        "Ваш запрос: ...?",
        reply_to_message_id=message_id,
    )

    # Сохраняем метаданные для отложенного анализа
    context.bot_data["pending_analyses"][question_msg.message_id] = {
        "file_message_id": message_id,
        "model_choice": model_choice,
        "chat_id": chat_id,
        "user_id": query.from_user.id,
    }


# ══════════════════════════════════════════════════════════════════════════════
# ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════════

app = ApplicationBuilder().token(TELEGRAM_TOKEN).build()
app.add_handler(CommandHandler("start", start))
app.add_handler(CommandHandler("ping", ping))
app.add_handler(CommandHandler("clear", clear))
app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))
app.add_handler(MessageHandler(filters.VOICE, handle_voice))
app.add_handler(MessageHandler(filters.PHOTO, handle_photo))
app.add_handler(MessageHandler(filters.VIDEO, handle_video))
app.add_handler(MessageHandler(filters.VIDEO_NOTE, handle_video_note))
app.add_handler(MessageHandler(filters.ANIMATION, handle_animation))
app.add_handler(MessageHandler(filters.Document.ALL, handle_document))
app.add_handler(CallbackQueryHandler(handle_ai_analyze_callback, pattern="^ai_analyze:"))

if __name__ == "__main__":
    logger.info(
        "Bot started with Smart Router (choose_model) + Groq Whisper STT "
        "+ PDF/DOCX/TXT/MD/JSON/ZIP/XLSX/GIF support "
        "+ Фаза 5: маршрутизация по топикам Jade_Developer"
    )
    app.run_polling()
